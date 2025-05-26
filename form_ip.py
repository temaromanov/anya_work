import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from tkcalendar import DateEntry
import pandas as pd
import os
from docx import Document

class ExcelEntryAppIP:
    def __init__(self, root):
        self.root = root
        self.root.title("Форма заполнения реквизитов (ИП)")
        self.root.geometry("760x500")
        self.entries = {}
        self.excel_file = None
        self.word_template = "Шаблон_аренда_договора.docx"  # Можешь поменять под ИП

        style = ttk.Style()
        style.configure("TNotebook.Tab", font=("Segoe UI", 10, "bold"))
        style.configure("TLabel", font=("Segoe UI", 10))
        style.configure("TButton", font=("Segoe UI", 10))

        notebook = ttk.Notebook(root)
        notebook.pack(expand=True, fill='both', padx=12, pady=12)

        self.frames = {
            "Основное": ttk.Frame(notebook, padding=10),
            "Банк и документы": ttk.Frame(notebook, padding=10),
            "Даты и номер": ttk.Frame(notebook, padding=10)
        }
        for name, frame in self.frames.items():
            notebook.add(frame, text=name)

        # Примерные поля — замени на свои!
        self.fields_main = [
            "ФИО", "ИНН", "ОГРНИП", "Паспорт", "Дата_рождения", "Адрес_регистрации", "номер_договора"
        ]
        self.fields_bank = [
            "Банк", "БИК", "к_счет", "Расч_счет"
        ]
        self.fields_date = [
            "Дата_договора", "Дата"
        ]
        bank_options = ["ПАО СБЕРБАНК", "ВТБ", "Газпромбанк", "Альфа-Банк", "Тинькофф"]

        for field in self.fields_main:
            self.create_row(self.frames["Основное"], field)

        for field in self.fields_bank:
            if field == "Банк":
                self.create_combobox(self.frames["Банк и документы"], field, bank_options)
            else:
                self.create_row(self.frames["Банк и документы"], field)

        for field in self.fields_date:
            if "Дата" in field:
                self.create_dateentry(self.frames["Даты и номер"], field)
            else:
                self.create_row(self.frames["Даты и номер"], field)

        button_frame = ttk.Frame(root)
        button_frame.pack(pady=15)

        ttk.Button(button_frame, text="📂 Выбрать Excel-файл", command=self.choose_file).pack(side=tk.LEFT, padx=8)
        ttk.Button(button_frame, text="💾 Сохранить как...", command=self.save_as).pack(side=tk.LEFT, padx=8)
        ttk.Button(button_frame, text="✅ Добавить и создать Word", command=self.save_and_generate_word).pack(side=tk.LEFT, padx=8)

    def create_row(self, parent, field):
        frame = ttk.Frame(parent)
        frame.pack(fill='x', pady=3)
        label = ttk.Label(frame, text=field.replace("_", " "), width=25, anchor='w')
        label.pack(side='left')
        entry = ttk.Entry(frame, width=65)
        entry.pack(side='left', padx=5)
        self.entries[field] = entry

    def create_combobox(self, parent, field, options):
        frame = ttk.Frame(parent)
        frame.pack(fill='x', pady=3)
        label = ttk.Label(frame, text=field.replace("_", " "), width=25, anchor='w')
        label.pack(side='left')
        combo = ttk.Combobox(frame, values=options, width=63)
        combo.pack(side='left', padx=5)
        self.entries[field] = combo

    def create_dateentry(self, parent, field):
        frame = ttk.Frame(parent)
        frame.pack(fill='x', pady=3)
        label = ttk.Label(frame, text=field.replace("_", " "), width=25, anchor='w')
        label.pack(side='left')
        date_entry = DateEntry(frame, width=61, date_pattern="dd.mm.yyyy")
        date_entry.pack(side='left', padx=5)
        self.entries[field] = date_entry

    def choose_file(self):
        path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if path:
            self.excel_file = path

    def save_as(self):
        path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
        if path:
            self.excel_file = path

    def save_and_generate_word(self):
        if not self.excel_file:
            messagebox.showwarning("Файл не выбран", "Пожалуйста, выберите или создайте Excel-файл перед сохранением.")
            return

        new_data = {k: v.get() for k, v in self.entries.items()}
        new_row = pd.DataFrame([new_data])

        if os.path.exists(self.excel_file):
            try:
                df = pd.read_excel(self.excel_file)
                df = pd.concat([df, new_row], ignore_index=True)
            except:
                df = new_row
        else:
            df = new_row

        try:
            df.to_excel(self.excel_file, index=False)
        except PermissionError:
            messagebox.showerror("Ошибка", f"Файл занят — закрой его в Excel:\n{self.excel_file}")
            return

        if os.path.exists(self.word_template):
            doc = Document(self.word_template)
            for p in doc.paragraphs:
                for run in p.runs:
                    for key, val in new_data.items():
                        if f"{{{{{key}}}}}" in run.text:
                            run.text = run.text.replace(f"{{{{{key}}}}}", val)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                for key, val in new_data.items():
                                    if f"{{{{{key}}}}}" in run.text:
                                        run.text = run.text.replace(f"{{{{{key}}}}}", val)
            word_output = self.excel_file.replace(".xlsx", "_документ.docx")
            doc.save(word_output)
            os.startfile(word_output)

        for entry in self.entries.values():
            entry.delete(0, tk.END)
