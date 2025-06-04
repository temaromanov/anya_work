
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from tkcalendar import DateEntry
import pandas as pd
import os
from docx import Document
from datetime import datetime
import pymorphy3

morph = pymorphy3.MorphAnalyzer()

def fio_to_rod_and_short(fio_nominative):
    parts = fio_nominative.strip().split()
    if len(parts) < 3:
        return fio_nominative, fio_nominative
    fam, name, otch = parts
    fam_rod = morph.parse(fam)[0].inflect({'gent'}).word.title()
    name_rod = morph.parse(name)[0].inflect({'gent'}).word.title()
    otch_rod = morph.parse(otch)[0].inflect({'gent'}).word.title()
    fio_rod = f"{fam_rod} {name_rod} {otch_rod}"
    fio_short = f"{fam} {name[0]}.{otch[0]}."
    return fio_rod, fio_short

def format_date_verbose(date_str):
    months = ['', 'января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
              'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря']
    try:
        dt = datetime.strptime(date_str, "%d.%m.%Y")
        return f"«{dt.day:02d}» {months[dt.month]} {dt.year} г."
    except Exception:
        return date_str

def replace_variables_in_doc(doc, replacements):
    for paragraph in doc.paragraphs:
        inline_replace_in_runs(paragraph.runs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    inline_replace_in_runs(paragraph.runs, replacements)

def inline_replace_in_runs(runs, replacements):
    full_text = "".join([run.text for run in runs])
    for key, value in replacements.items():
        full_text = full_text.replace(f"{{{{{key}}}}}", str(value))
    idx = 0
    for run in runs:
        run_len = len(run.text)
        run.text = full_text[idx:idx+run_len]
        idx += run_len

class ExcelEntryAppActVozvrataOOO:
    def __init__(self, parent, go_back):
        self.parent = parent
        self.go_back_callback = go_back
        self.entries = {}
        self.excel_file = None
        self.word_template = "Шаблон_акт_возврата_ООО.docx"

        btn_back = ttk.Button(self.parent, text="← Назад к выбору", command=self.go_back)
        label = ttk.Label(parent, text="Акт возврата (ООО)", font=("Segoe UI", 17, "bold"))
        label.pack(pady=(10, 10))
        btn_back.pack(anchor="nw", padx=10, pady=10)

        style = ttk.Style()
        style.configure("TNotebook.Tab", font=("Segoe UI", 10, "bold"))
        style.configure("TLabel", font=("Segoe UI", 10))
        style.configure("TButton", font=("Segoe UI", 10))

        notebook = ttk.Notebook(self.parent)
        notebook.pack(expand=True, fill='both', padx=12, pady=12)

        self.frame = ttk.Frame(notebook, padding=10)
        notebook.add(self.frame, text="Данные для акта")

        self.fields = [
            "ФИО_им", "Должность", "Номер_договора", "Юридический_адрес", "Фактический_адрес",
            "ИНН", "ОГРН", "КПП", "ОКПО", "Расч_счет", "Банк", "БИК", "к_счет",
            "Дата", "Название_ООО"
        ]

        bank_options = ["ПАО СБЕРБАНК", "ВТБ", "Газпромбанк", "Альфа-Банк", "Тинькофф"]
        person_options = ["Генеральный директор", "Президент", "Директор"]

        for field in self.fields:
            if field == "ФИО_им":
                self.create_row(self.frame, field, "ФИО")
            elif field == "Должность":
                self.create_combobox(self.frame, field, person_options)
            elif field == "Банк":
                self.create_combobox(self.frame, field, bank_options)
            elif "Дата" in field:
                self.create_dateentry(self.frame, field)
            else:
                self.create_row(self.frame, field)

        button_frame = ttk.Frame(self.parent)
        button_frame.pack(pady=15)

        ttk.Button(button_frame, text="📂 Выбрать Excel-файл", command=self.choose_file).pack(side=tk.LEFT, padx=8)
        ttk.Button(button_frame, text="💾 Сохранить как...", command=self.save_as).pack(side=tk.LEFT, padx=8)
        ttk.Button(button_frame, text="✅ Добавить и создать Word", command=self.save_and_generate_word).pack(side=tk.LEFT, padx=8)

    def go_back(self):
        self.parent.pack_forget()
        self.go_back_callback()

    def create_row(self, parent, field, label_text=None):
        frame = ttk.Frame(parent)
        frame.pack(fill='x', pady=3)
        label = ttk.Label(frame, text=label_text or field.replace("_", " "), width=40, anchor='w')
        label.pack(side='left')
        entry = ttk.Entry(frame, width=55)
        entry.pack(side='left', padx=5)
        self.entries[field] = entry

    def create_combobox(self, parent, field, options):
        frame = ttk.Frame(parent)
        frame.pack(fill='x', pady=3)
        label = ttk.Label(frame, text=field.replace("_", " "), width=25, anchor='w')
        label.pack(side='left')
        combo = ttk.Combobox(frame, values=options, width=63)
        combo.pack(side='left', padx=5)
        self.entries[field] = combo

    def create_dateentry(self, parent, field):
        frame = ttk.Frame(parent)
        frame.pack(fill='x', pady=3)
        label = ttk.Label(frame, text=field.replace("_", " "), width=25, anchor='w')
        label.pack(side='left')
        date_entry = DateEntry(frame, width=61, date_pattern="dd.mm.yyyy")
        date_entry.pack(side='left', padx=5)
        self.entries[field] = date_entry

    def choose_file(self):
        path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if path:
            self.excel_file = path

    def save_as(self):
        path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
        if path:
            self.excel_file = path

    def save_and_generate_word(self):
        if not self.excel_file:
            messagebox.showwarning("Файл не выбран", "Пожалуйста, выберите или создайте Excel-файл перед сохранением.")
            return

        fio_input = self.entries["ФИО_им"].get()
        fio_rod, fio_short = fio_to_rod_and_short(fio_input)

        new_data = {k: v.get() for k, v in self.entries.items()}
        new_data["Полное_имя_род_падеж"] = fio_rod
        new_data["Сокрщ_имя_дир"] = fio_short

        date_str = new_data.get("Дата", "")
        new_data["Дата_прописью"] = format_date_verbose(date_str) if date_str else ""

        new_row = pd.DataFrame([new_data])
        if os.path.exists(self.excel_file):
            try:
                df = pd.read_excel(self.excel_file)
                df = pd.concat([df, new_row], ignore_index=True)
            except:
                df = new_row
        else:
            df = new_row

        try:
            df.to_excel(self.excel_file, index=False)
        except PermissionError:
            messagebox.showerror("Ошибка", f"Файл занят — закрой его в Excel:{self.excel_file}")
            return

        if os.path.exists(self.word_template):
            doc = Document(self.word_template)
            replace_variables_in_doc(doc, new_data)
            word_output = self.excel_file.replace(".xlsx", "_акт_возврата.docx")
            doc.save(word_output)
            os.startfile(word_output)
        else:
            messagebox.showerror("Нет шаблона", f"Word-шаблон не найден:{self.word_template}")

        for entry in self.entries.values():
            entry.delete(0, tk.END)
