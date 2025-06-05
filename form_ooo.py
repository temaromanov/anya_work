import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from tkcalendar import DateEntry
import pandas as pd
import os
from docx import Document
import pymorphy3
from num2words import num2words
from datetime import datetime

def copy_paste_fix(e):
        if e.state & 0x4:  # Ctrl зажат
            if e.keycode == 86:  # V
                e.widget.event_generate('<<Paste>>')
            elif e.keycode == 67:  # C
                e.widget.event_generate('<<Copy>>')
            elif e.keycode == 88:  # X
                e.widget.event_generate('<<Cut>>')

morph = pymorphy3.MorphAnalyzer()

def fio_to_rod_and_short(fio_nominative):
    parts = fio_nominative.strip().split()
    if len(parts) < 3:
        return fio_nominative, fio_nominative

    fam, name, otch = parts

    fam_parsed = morph.parse(fam)[0].inflect({'gent'})
    name_parsed = morph.parse(name)[0].inflect({'gent'})
    otch_parsed = morph.parse(otch)[0].inflect({'gent'})

    fam_rod = fam_parsed.word.title() if fam_parsed else fam
    name_rod = name_parsed.word.title() if name_parsed else name
    otch_rod = otch_parsed.word.title() if otch_parsed else otch

    fio_rod = f"{fam_rod} {name_rod} {otch_rod}"
    fio_short = f"{fam} {name[0]}.{otch[0]}."

    return fio_rod, fio_short

def rub_to_words(rub):
    rub = int(rub)
    rub_text = num2words(rub, lang="ru")
    return f"({rub_text})"

def format_date_verbose(date_str):
    months = [
        '', 'января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
        'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря'
    ]
    try:
        dt = datetime.strptime(date_str, "%d.%m.%Y")
        return f"«{dt.day:02d}» {months[dt.month]} {dt.year} г."
    except Exception:
        return date_str

def replace_variables_in_doc(doc, replacements):
    def replace_in_paragraph(paragraph):
        for key, value in replacements.items():
            if f"{{{{{key}}}}}" in paragraph.text:
                inline_replace = paragraph.text.replace(f"{{{{{key}}}}}", str(value))
                paragraph.clear()
                paragraph.add_run(inline_replace)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)

def inline_replace_in_runs(runs, replacements):
    full_text = "".join([run.text for run in runs])
    for key, value in replacements.items():
        full_text = full_text.replace(f"{{{{{key}}}}}", str(value))
    idx = 0
    for run in runs:
        run_len = len(run.text)
        run.text = full_text[idx:idx+run_len]
        idx += run_len

class ExcelEntryAppOOO:
    def __init__(self, parent, go_back):
        self.parent = parent
        self.go_back_callback = go_back
        self.entries = {}
        self.excel_file = None
        self.word_template = "Шаблон_аренда_договор_ооо.docx"

        btn_back = ttk.Button(self.parent, text="← Назад к выбору", command=self.go_back)
        label = ttk.Label(parent, text="Договор аренды (ООО)", font=("Segoe UI", 17, "bold"))
        label.pack(pady=(10, 10))
        btn_back.pack(anchor="nw", padx=10, pady=10)

        style = ttk.Style()
        style.configure("TNotebook.Tab", font=("Segoe UI", 10, "bold"))
        style.configure("TLabel", font=("Segoe UI", 10))
        style.configure("TButton", font=("Segoe UI", 10))

        notebook = ttk.Notebook(self.parent)
        notebook.pack(expand=True, fill='both', padx=12, pady=12)

        self.frames = {
            "Основное": ttk.Frame(notebook, padding=10),
            "Банк и документы": ttk.Frame(notebook, padding=10),
            "Даты и номер": ttk.Frame(notebook, padding=10)
        }

        for name, frame in self.frames.items():
            notebook.add(frame, text=name)

        self.fields_main = [
            "Название_организации",
            "Должность",
            "ФИО_им",
            "Мероприятие", "Сумма_арендной_платы"
        ]
        self.fields_bank = [
            "Юридический_адрес", "Фактический_адрес",
            "ИНН", "КПП", "ОКПО", "ОГРН", "Банк", "БИК", "Расч_счет", "к_счет"
        ]
        self.fields_date = [
            "Номер_договора", "Дата_начала_аренды", "Дата_окончания_аренды", "Дата"
        ]

        bank_options = ["ПАО СБЕРБАНК", "ВТБ", "Газпромбанк", "Альфа-Банк", "Тинькофф"]
        person_options = ["Генеральный директор", "Президент", "Директор"]

        self.nds_label = None

        for field in self.fields_main:
            if field == "ФИО_им":
                self.create_row(self.frames["Основное"], field, "ФИО")
            elif field == "Должность":
                self.create_combobox(self.frames["Основное"], field, person_options)
            elif field == "Сумма_арендной_платы":
                self.create_row(self.frames["Основное"], field)
                self.nds_label = ttk.Label(self.frames["Основное"], text="НДС (5%): 0.00", font=("Segoe UI", 10, "bold"), foreground="#225500")
                self.nds_label.pack(anchor="w", padx=10, pady=2)
                self.entries[field].bind("<KeyRelease>", self.update_nds)
            else:
                self.create_row(self.frames["Основное"], field)

        for field in self.fields_bank:
            if field == "Контактная_персона":
                self.create_combobox(self.frames["Банк и документы"], field, person_options)
            else:
                self.create_row(self.frames["Банк и документы"], field)

        for field in self.fields_date:
            if "Дата" in field:
                self.create_dateentry(self.frames["Даты и номер"], field)
            else:
                self.create_row(self.frames["Даты и номер"], field)

        button_frame = ttk.Frame(self.parent)
        button_frame.pack(pady=15)

        ttk.Button(button_frame, text="📂 Выбрать Excel-файл", command=self.choose_file).pack(side=tk.LEFT, padx=8)
        ttk.Button(button_frame, text="💾 Сохранить как...", command=self.save_as).pack(side=tk.LEFT, padx=8)
        ttk.Button(button_frame, text="✅ Добавить и создать Word", command=self.save_and_generate_word).pack(side=tk.LEFT, padx=8)

        # Обрабатываем Ctrl+C/V/X даже при русской раскладке
        self.parent.bind_all("<Key>", copy_paste_fix)


    def go_back(self):
        self.parent.pack_forget()
        self.go_back_callback()

    def create_row(self, parent, field, label_text=None):
        frame = ttk.Frame(parent)
        frame.pack(fill='x', pady=3)
        label = ttk.Label(frame, text=label_text or field.replace("_", " "), width=40, anchor='w')
        label.pack(side='left')
        entry = ttk.Entry(frame, width=55)
        entry.pack(side='left', padx=5)
        self.entries[field] = entry

    def create_combobox(self, parent, field, options):
        frame = ttk.Frame(parent)
        frame.pack(fill='x', pady=3)
        label = ttk.Label(frame, text=field.replace("_", " "), width=25, anchor='w')
        label.pack(side='left')
        combo = ttk.Combobox(frame, values=options, width=63)
        combo.pack(side='left', padx=5)
        self.entries[field] = combo

    def create_dateentry(self, parent, field):
        frame = ttk.Frame(parent)
        frame.pack(fill='x', pady=3)
        label = ttk.Label(frame, text=field.replace("_", " "), width=25, anchor='w')
        label.pack(side='left')
        date_entry = DateEntry(frame, width=61, date_pattern="dd.mm.yyyy")
        date_entry.pack(side='left', padx=5)
        self.entries[field] = date_entry

    def choose_file(self):
        path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if path:
            self.excel_file = path

    def save_as(self):
        path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
        if path:
            self.excel_file = path

    def update_nds(self, event=None):
        try:
            summa = float(self.entries["Сумма_арендной_платы"].get().replace(",", "."))
            nds = round(summa * 5 / 105, 2)
            if self.nds_label:
                self.nds_label.config(text=f"НДС (5%): {nds:.2f}")
        except Exception:
            if self.nds_label:
                self.nds_label.config(text="НДС (5%): 0.00")

    def save_and_generate_word(self):
        if not self.excel_file:
            messagebox.showwarning("Файл не выбран", "Пожалуйста, выберите или создайте Excel-файл перед сохранением.")
            return

        fio_input = self.entries["ФИО_им"].get()
        fio_rod, fio_short = fio_to_rod_and_short(fio_input)

        new_data = {k: v.get() for k, v in self.entries.items()}
        new_data["Полное_имя_род_падеж"] = fio_rod
        new_data["Сокрщ_имя_дир"] = fio_short

        position = new_data.get("Должность", "")
        if position:
            words = position.strip().split()
            inflected_words = []
            for word in words:
                parsed = morph.parse(word)[0]
                inflected = parsed.inflect({'gent'})
                inflected_words.append(inflected.word if inflected else word)
            if inflected_words:
                inflected_words[0] = inflected_words[0].capitalize()
            new_data["Должность_род"] = " ".join(inflected_words)
        else:
            new_data["Должность_род"] = ""

        summa_str = new_data.get("Сумма_арендной_платы", "0").replace(",", ".")
        if summa_str.strip() == "":
            summa_str = "0"
        try:
            summa_float = float(summa_str)
        except Exception:
            summa_float = 0

        rub, sep, kop = summa_str.partition(".")
        if not sep:
            rub, kop = rub, "00"
        else:
            kop = (kop + "00")[:2]
        new_data["Сумма_арендной_платы_руб"] = rub if rub else "0"
        new_data["Сумма_арендной_платы_коп"] = kop if kop else "00"
        new_data["Сумма_арендной_платы_прописью"] = rub_to_words(rub if rub else "0")

        nds = round(summa_float * 5 / 105, 2) if summa_float else 0
        new_data["НДС"] = f"{nds:.2f}"

        # --- ДАТА ПРОПИСЬЮ ---
        date_str = new_data.get("Дата", "")
        if date_str:
            new_data["Дата_прописью"] = format_date_verbose(date_str)
        else:
            new_data["Дата_прописью"] = ""
        # --- ДАТА ОКОНЧАНИЯ ПРОПИСЬЮ ---
        date_end_str = new_data.get("Дата_окончания_аренды", "")
        if date_end_str:
            new_data["Дата_окончания_прописью"] = format_date_verbose(date_end_str)
        else:
            new_data["Дата_окончания_прописью"] = ""

        new_row = pd.DataFrame([new_data])

        if os.path.exists(self.excel_file):
            try:
                df = pd.read_excel(self.excel_file)
                df = pd.concat([df, new_row], ignore_index=True)
            except:
                df = new_row
        else:
            df = new_row

        try:
            df.to_excel(self.excel_file, index=False)
        except PermissionError:
            messagebox.showerror("Ошибка", f"Файл занят — закрой его в Excel:\n{self.excel_file}")
            return

        # --- УНИВЕРСАЛЬНАЯ ЗАМЕНА В WORD ---
        if os.path.exists(self.word_template):
            doc = Document(self.word_template)
            replace_variables_in_doc(doc, new_data)
            word_output = self.excel_file.replace(".xlsx", "_документ.docx")
            doc.save(word_output)
            os.startfile(word_output)
        else:
            messagebox.showerror("Нет шаблона", f"Word-шаблон не найден:\n{self.word_template}")

        for entry in self.entries.values():
            entry.delete(0, tk.END)
        self.update_nds()
